import os
import re
import logging
from pathlib import Path
from .skills import additional_skills
import spacy
from datetime import datetime
import dateutil.parser

# Set up logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Load spaCy model
try:
    nlp = spacy.load("en_core_web_sm")
    SPACY_AVAILABLE = True
except (ImportError, OSError):
    logger.warning("spaCy model not available. Install with: python -m spacy download en_core_web_sm")
    SPACY_AVAILABLE = False

# Try to import document processing libraries
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    logger.warning("PyPDF2 not available. PDF parsing will be limited.")
    PDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    logger.warning("python-docx not available. DOCX parsing will be limited.")
    DOCX_AVAILABLE = False

def extract_text_from_pdf(pdf_path):
    """Extract text from PDF file"""
    if not PDF_AVAILABLE:
        return "PDF parsing not available. Install PyPDF2."
    
    text = ""
    try:
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        text = f"Error extracting text: {str(e)}"
    
    return text

def extract_text_from_docx(docx_path):
    """Extract text from DOCX file"""
    if not DOCX_AVAILABLE:
        return "DOCX parsing not available. Install python-docx."
    
    text = ""
    try:
        doc = docx.Document(docx_path)
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Each cell may contain paragraphs
                    for paragraph in cell.paragraphs:
                        text += paragraph.text + "\n"
        
        # Extract text from headers
        for section in doc.sections:
            for header in [section.header, section.footer]:
                if header:
                    for para in header.paragraphs:
                        text += para.text + "\n"
                    
                    # Headers can also contain tables
                    for table in header.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    text += paragraph.text + "\n"
        
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        text = f"Error extracting text: {str(e)}"
    
    return text

def extract_text(file_path):
    """Extract text from a file based on its extension"""
    file_extension = Path(file_path).suffix.lower()
    
    if file_extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    else:
        return f"Unsupported file format: {file_extension}"

def extract_email(text):
    """Extract email addresses from text using spaCy and regex"""
    # Try with spaCy first if available
    if SPACY_AVAILABLE:
        doc = nlp(text)
        for token in doc:
            if token.like_email:
                return token.text
    
    # Fallback to regex with improved pattern
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    emails = re.findall(email_pattern, text)
    return emails[0] if emails else None

def extract_phone(text):
    """Extract phone numbers from text using improved patterns"""
    # Handle multiple phone formats
    # Basic patterns
    phone_patterns = [
        # Standard formats
        r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', 
        r'(?:\+?\d{1,3}[-.\s]?)?\d{3}[-.\s]?\d{3}[-.\s]?\d{4}',        
        r'(?:\+?\d{1,3}[-.\s]?)?\d{10}',                              
        
        # International formats
        r'\+\d{1,3}[-.\s]?\d{1,14}',                                   
        
        # With extensions
        r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}(?:[-.\s]?(?:ext|x|ext.)[-.\s]?\d{1,5})?',
        
        # Format with periods
        r'\d{3}\.\d{3}\.\d{4}'                                          
    ]
    
    # Try each pattern
    for pattern in phone_patterns:
        matches = re.findall(pattern, text)
        if matches:
            # Clean up the first match
            phone = matches[0].strip()
            # Standardize formatting
            phone = re.sub(r'\s+', ' ', phone)  
            return phone
    
    # If spaCy is available, try to use it as a backup
    if SPACY_AVAILABLE:
        doc = nlp(text)
        for ent in doc.ents:
            if ent.label_ == "CARDINAL" and len(ent.text) >= 7:
                digits = sum(c.isdigit() for c in ent.text)
                if digits >= 7:
                    return ent.text
    
    return None

def extract_skills(text):
    """Extract skills from text using spaCy and the skills list"""
    # Common tech skills 
    common_skills = [
        "python", "java", "javascript", "html", "css", "react", "angular", "vue", "node", 
        "express", "flask", "django", "spring", "sql", "nosql", "mongodb", "mysql", 
        "postgresql", "aws", "azure", "gcp", "docker", "kubernetes", "git", "jenkins",
        "ci/cd", "machine learning", "data science", "artificial intelligence", "nlp"
    ] + additional_skills
    
    skills = []
    text_lower = text.lower()
    
    # Extract skills from the text using regex with word boundaries
    for skill in common_skills:
        # Use word boundaries to find complete words
        if re.search(r'\b' + re.escape(skill) + r'\b', text_lower):
            skills.append(skill)
    
    # If spaCy is available, use it to extract more skills
    if SPACY_AVAILABLE:
        doc = nlp(text)
        
        # Look for noun phrases that might be skills
        for chunk in doc.noun_chunks:
            chunk_text = chunk.text.lower()
            # Check if this noun phrase contains any of our skills
            for skill in common_skills:
                if skill in chunk_text and skill not in skills:
                    skills.append(skill)
        
        # Look for technical terms using entity recognition
        for ent in doc.ents:
            if ent.label_ == "ORG" or ent.label_ == "PRODUCT":
                ent_text = ent.text.lower()
                if any(skill in ent_text for skill in common_skills) and ent_text not in skills:
                    skills.append(ent_text)
    
    return sorted(list(set(skills)))  # Remove duplicates and sort
def extract_certifications(text):
    """
    Enhanced function to extract certification information from text using NLP and regex
    
    Returns a list of dictionaries with certification details:
    [{"name": "cert name", "provider": "issuing org", "cert_number": "ABC123", "date": "Jan 2022"}]
    """
    certifications = []
    
    # Split text into lines for processing
    lines = text.split('\n')
    
    # Find certification section
    cert_section_start = None
    cert_section_end = None
    
    # Look for certification section header
    for i, line in enumerate(lines):
        if re.search(r'^(?:certifications?|qualifications?|credentials?|professional\s+certifications?)\s*:?', 
                    line, re.IGNORECASE):
            cert_section_start = i + 1
            break
    
    # If certification section found, find its end
    if cert_section_start is not None:
        for i in range(cert_section_start, len(lines)):
            if re.search(r'^(?:education|experience|skills|references|projects|awards|publications|languages|interests|hobbies)', 
                        lines[i], re.IGNORECASE):
                cert_section_end = i
                break
        
        # If no end found, assume it goes to the end of the document
        if cert_section_end is None:
            cert_section_end = len(lines)
    
    # Extract certification text if section found
    cert_text = ""
    if cert_section_start is not None and cert_section_end is not None:
        cert_text = '\n'.join(lines[cert_section_start:cert_section_end])
    else:
        # If no clear section, use the entire text
        cert_text = text
    
    # Pattern 1: Standard certification format
    # Example: "AWS Certified Solutions Architect - Professional (AWS-CSA-P) - Amazon Web Services, June 2022"
    std_cert_pattern = r"([\w\s\-\(\)]+?)(?:\s*[-–:]\s*|\s+by\s+|\s+from\s+)([\w\s\-\&\.]+)(?:,\s*|\s+|\s*\()((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|(?:\d{1,2}/\d{1,2}/\d{4}|\d{4}))"
    
    # Pattern 2: Certification with numbers
    # Example: "Certification No. ABC123 - Microsoft Certified: Azure Administrator Associate"
    cert_num_pattern = r"(?:certification|cert|certificate|license)(?:\s+(?:no|number|#|id|code))(?:\s*[:.]\s*|\s+)([\w\-]+)\s*(?:\)|,|\.|\s+)\s*(.*?)(?:\s+by\s+|\s+from\s+|\s*[-–:]\s*)([\w\s\-\&\.]+)"
    
    # Pattern 3: Simple certification listing
    # Example: "CISSP (Certified Information Systems Security Professional)"
    simple_cert_pattern = r"(?:•|\*|\-|\d+\.\s*)([\w\s\-\(\):]+?)(?:\s*[-–:]\s*|\s+by\s+|\s+from\s+|\s+issued\s+by\s+|\s*$)([\w\s\-\&\.]*)"
    
    # Process the text with spaCy if available
    cert_entities = []
    if SPACY_AVAILABLE:
        doc = nlp(cert_text)
        
        # Extract potential certification names using entity recognition
        for ent in doc.ents:
            if ent.label_ in ["ORG", "PRODUCT", "WORK_OF_ART"]:
                cert_entities.append(ent.text)
        
        # Extract potential dates near certifications for expiration/issue dates
        dates = [ent.text for ent in doc.ents if ent.label_ == "DATE"]
    
    # Apply standard certification pattern
    for match in re.finditer(std_cert_pattern, cert_text, re.IGNORECASE):
        cert = {
            "name": match.group(1).strip(),
            "provider": match.group(2).strip(),
            "cert_number": "",
            "date": match.group(3).strip()
        }
        
        # Check for certification number in the name
        num_match = re.search(r'\(([A-Z0-9\-]+)\)', cert["name"])
        if num_match:
            potential_num = num_match.group(1).strip()
            # Verify it looks like a certification number (contains numbers or specific patterns)
            if re.search(r'\d', potential_num) or re.search(r'^[A-Z]+-[A-Z0-9]+$', potential_num):
                cert["cert_number"] = potential_num
                # Remove the number from the name
                cert["name"] = re.sub(r'\s*\([A-Z0-9\-]+\)', '', cert["name"]).strip()
        
        certifications.append(cert)
    
    # Apply certification number pattern
    for match in re.finditer(cert_num_pattern, cert_text, re.IGNORECASE):
        cert = {
            "name": match.group(2).strip(),
            "provider": match.group(3).strip() if len(match.groups()) >= 3 else "",
            "cert_number": match.group(1).strip(),
            "date": ""
        }
        
        # Try to find a nearby date
        if SPACY_AVAILABLE:
            # Get position of the match
            start_pos = match.start()
            end_pos = match.end()
            
            # Look for date entities within reasonable proximity
            doc = nlp(cert_text[max(0, start_pos - 100):min(len(cert_text), end_pos + 100)])
            for ent in doc.ents:
                if ent.label_ == "DATE":
                    cert["date"] = ent.text
                    break
        
        certifications.append(cert)
    
    # Apply simple certification pattern
    for match in re.finditer(simple_cert_pattern, cert_text, re.IGNORECASE):
        # Skip if it's too short or likely not a certification
        if len(match.group(1).strip()) < 3:
            continue
            
        cert = {
            "name": match.group(1).strip(),
            "provider": match.group(2).strip() if len(match.groups()) >= 2 and match.group(2) else "",
            "cert_number": "",
            "date": ""
        }
        
        # Handle certification acronyms in parentheses or vice versa
        name = cert["name"]
        acronym_match1 = re.search(r'(.*?)\s*\(([A-Z]{2,})\)', name)  # Name (ACRONYM)
        acronym_match2 = re.search(r'([A-Z]{2,})\s*\((.*?)\)', name)  # ACRONYM (Name)
        
        if acronym_match1:
            cert["name"] = f"{acronym_match1.group(1).strip()} ({acronym_match1.group(2)})"
        elif acronym_match2:
            cert["name"] = f"{acronym_match2.group(1)} ({acronym_match2.group(2).strip()})"
        
        # Try to find a nearby date if spaCy is available
        if SPACY_AVAILABLE:
            # Get position of the match
            start_pos = match.start()
            end_pos = match.end()
            
            # Define search context (100 chars before and after)
            context_start = max(0, start_pos - 100)
            context_end = min(len(cert_text), end_pos + 100)
            
            # Look for date entities within proximity
            context_doc = nlp(cert_text[context_start:context_end])
            for ent in context_doc.ents:
                if ent.label_ == "DATE":
                    cert["date"] = ent.text
                    break
        
        certifications.append(cert)
    
    # Process certification entities found by spaCy if still no certifications
    if not certifications and cert_entities:
        for entity in cert_entities:
            # Skip if it's too short
            if len(entity) < 3:
                continue
                
            cert = {
                "name": entity,
                "provider": "",
                "cert_number": "",
                "date": ""
            }
            
            # Try to find provider and date in nearby text
            if SPACY_AVAILABLE:
                # Find the entity in the text
                entity_pos = cert_text.find(entity)
                if entity_pos != -1:
                    # Get surrounding context
                    context_start = max(0, entity_pos - 100)
                    context_end = min(len(cert_text), entity_pos + entity.__len__() + 100)
                    context = cert_text[context_start:context_end]
                    
                    # Look for organizations that might be providers
                    context_doc = nlp(context)
                    for ent in context_doc.ents:
                        if ent.label_ == "ORG" and ent.text != entity:
                            cert["provider"] = ent.text
                            break
                    
                    # Look for dates
                    for ent in context_doc.ents:
                        if ent.label_ == "DATE":
                            cert["date"] = ent.text
                            break
            
            certifications.append(cert)
    
    # Look for list patterns in certification section (if found)
    if cert_section_start is not None and cert_section_end is not None:
        section_lines = lines[cert_section_start:cert_section_end]
        current_cert = None
        
        for line in section_lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if this line starts a new certification entry
            # Indicated by bullet point, number, or just a new line with capitalization
            is_new_entry = (re.match(r'^(?:•|\*|\-|\d+\.|\–)\s+', line) or 
                          (line[0].isupper() if line else False))
            
            if is_new_entry or current_cert is None:
                # If we have a previous cert, save it
                if current_cert is not None:
                    # Only add if we don't already have it
                    if not any(c["name"] == current_cert["name"] for c in certifications):
                        certifications.append(current_cert)
                
                # Start new cert entry
                current_cert = {"name": line, "provider": "", "cert_number": "", "date": ""}
                
                # Clean up the bullet point or numbering
                current_cert["name"] = re.sub(r'^(?:•|\*|\-|\d+\.|\–)\s+', '', current_cert["name"])
                
                # Try to extract provider if in same line
                provider_match = re.search(r'(?:by|from|through|via)\s+([\w\s\-\&\.]+)(?:,|\s+|$)', current_cert["name"])
                if provider_match:
                    current_cert["provider"] = provider_match.group(1).strip()
                    # Remove provider from name
                    current_cert["name"] = current_cert["name"].replace(provider_match.group(0), "").strip()
                
                # Try to extract date if in same line
                date_match = re.search(r'(?:(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|(?:\d{1,2}/\d{1,2}/\d{4}|\d{4}))', current_cert["name"])
                if date_match:
                    current_cert["date"] = date_match.group(0).strip()
                    # Remove date from name
                    current_cert["name"] = current_cert["name"].replace(date_match.group(0), "").strip()
                    # Also remove any commas or parenthesis that might be left
                    current_cert["name"] = re.sub(r'(?:\s*,\s*|\s*\(\s*|\s*\)\s*)$', '', current_cert["name"]).strip()
            else:
                # Add this line to the current cert's details
                # Check for provider
                provider_match = re.search(r'(?:by|from|through|via|issued by)\s+([\w\s\-\&\.]+)', line)
                if provider_match and not current_cert["provider"]:
                    current_cert["provider"] = provider_match.group(1).strip()
                
                # Check for cert number
                num_match = re.search(r'(?:certification|cert|certificate|license)(?:\s+(?:no|number|#|id|code))(?:\s*[:.]\s*|\s+)([\w\-]+)', line, re.IGNORECASE)
                if num_match and not current_cert["cert_number"]:
                    current_cert["cert_number"] = num_match.group(1).strip()
                
                # Check for date
                date_match = re.search(r'(?:(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|(?:\d{1,2}/\d{1,2}/\d{4}|\d{4}))', line)
                if date_match and not current_cert["date"]:
                    current_cert["date"] = date_match.group(0).strip()
        
        # Add last certification if there is one
        if current_cert is not None:
            if not any(c["name"] == current_cert["name"] for c in certifications):
                certifications.append(current_cert)
    
    # Deduplicate certifications
    unique_certs = []
    seen_names = set()
    
    for cert in certifications:
        # Normalize name for comparison
        norm_name = re.sub(r'\s+', ' ', cert["name"].lower().strip())
        if norm_name and norm_name not in seen_names:
            seen_names.add(norm_name)
            unique_certs.append(cert)
    
    return unique_certs

def extract_work_experience(text):
    """
    Enhanced function to extract work experience information from text using NLP and regex
    
    Returns a list of dictionaries with work experience details:
    [{"title": "job title", "company": "company name", "duration": "date range", 
      "location": "job location", "description": "job description"}]
    """
    work_experiences = []
    
    # Define section headers that might indicate work experience
    experience_headers = [
        r'(?:professional\s+)?(?:work\s+)?experience',
        r'employment\s+history',
        r'work\s+history',
        r'career\s+history',
        r'professional\s+background',
        r'job\s+history',
        r'professional\s+experience',
        r'relevant\s+experience'
    ]
    
    # Date patterns for matching
    date_patterns = [
        # Month Year - Month Year or Present
        r'(?:(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}\s*[-–—~to]*\s*(?:(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|Present|Current|Now|Ongoing))',
        # MM/YYYY - MM/YYYY or Present
        r'(?:\d{1,2}/\d{4}\s*[-–—~to]*\s*(?:\d{1,2}/\d{4}|Present|Current|Now|Ongoing))',
        # YYYY - YYYY or Present
        r'(?:\d{4}\s*[-–—~to]*\s*(?:\d{4}|Present|Current|Now|Ongoing))',
        # MM/DD/YYYY - MM/DD/YYYY or Present
        r'(?:\d{1,2}/\d{1,2}/\d{4}\s*[-–—~to]*\s*(?:\d{1,2}/\d{1,2}/\d{4}|Present|Current|Now|Ongoing))'
    ]
    
    # Combined date pattern
    date_pattern = '|'.join(f'({pattern})' for pattern in date_patterns)
    
    # Split text into lines for processing
    lines = text.split('\n')
    
    # Find the experience section
    experience_start_idx = None
    experience_end_idx = None
    
    for i, line in enumerate(lines):
        # Check if this line is an experience header
        if any(re.search(pattern, line, re.IGNORECASE) for pattern in experience_headers):
            experience_start_idx = i + 1
            break
    
    # If we found the experience section, find where it ends
    if experience_start_idx is not None:
        for i in range(experience_start_idx, len(lines)):
            # Check if this line is the start of another section
            if re.search(r'^(?:education|skills|certifications|references|projects|awards|publications|languages|interests|hobbies|additional|profile|summary)', 
                         lines[i], re.IGNORECASE) and not re.search(r'experience', lines[i], re.IGNORECASE):
                experience_end_idx = i
                break
        
        # If we didn't find an end, assume it goes to the end of the document
        if experience_end_idx is None:
            experience_end_idx = len(lines)
        
        # Extract the experience text
        experience_text = '\n'.join(lines[experience_start_idx:experience_end_idx])
    else:
        # If we didn't find a clear experience section, use the whole text
        experience_text = text
    
    # Use spaCy for enhanced extraction if available
    if SPACY_AVAILABLE:
        doc = nlp(experience_text)
        
        # Find paragraphs or blocks that might represent job entries
        paragraphs = re.split(r'\n\s*\n', experience_text)
        
        for paragraph in paragraphs:
            # Skip if paragraph is too short
            if len(paragraph.strip()) < 15:
                continue
            
            # Initialize job entry
            job = {
                "title": "",
                "company": "",
                "duration": "",
                "location": "",
                "description": paragraph.strip()
            }
            
            # Extract dates
            date_matches = re.finditer(date_pattern, paragraph, re.IGNORECASE)
            for match in date_matches:
                job["duration"] = match.group(0).strip()
                # Remove duration from description
                job["description"] = job["description"].replace(match.group(0), "").strip()
                break  # Just use the first date range found
                
            # Process with spaCy for entity extraction
            para_doc = nlp(paragraph)
            
            # Extract companies (ORG entities)
            for ent in para_doc.ents:
                if ent.label_ == "ORG" and not job["company"]:
                    # Verify it's likely a company (not a university, etc.)
                    if not re.search(r'(?:university|college|school|institute|academy)', ent.text, re.IGNORECASE):
                        job["company"] = ent.text.strip()
                        # Remove company from description
                        job["description"] = job["description"].replace(ent.text, "").strip()
                
                # Extract locations (GPE or LOC entities)
                elif ent.label_ in ["GPE", "LOC"] and not job["location"]:
                    job["location"] = ent.text.strip()
                    # Remove location from description if it's clearly marked
                    loc_pattern = f"(?:in|at|from)\\s+{re.escape(ent.text)}"
                    job["description"] = re.sub(loc_pattern, "", job["description"]).strip()
            
            # Extract job title using patterns if not already found
            # Common job title patterns
            title_patterns = [
                r'(?:^|\n)([A-Z][a-zA-Z\s]+(?:Engineer|Developer|Manager|Director|Analyst|Designer|Consultant|Specialist|Officer|Administrator|Coordinator|Lead|Head|Chief|Architect))',
                r'(?:^|\n)([A-Z][a-zA-Z\s]+(?:Intern|Associate|Assistant|Representative|Supervisor|Technician|Agent|Coach|Tutor|Instructor))',
                r'(?:as\s+(?:a|an)\s+([A-Za-z\s]+(?:Engineer|Developer|Manager|Director|Analyst|Designer|Consultant|Specialist)))',
                r'(?:position\s+of\s+([A-Za-z\s]+(?:Engineer|Developer|Manager|Director|Analyst|Designer|Consultant|Specialist)))'
            ]
            
            for pattern in title_patterns:
                title_match = re.search(pattern, paragraph)
                if title_match and not job["title"]:
                    job["title"] = title_match.group(1).strip()
                    break
            
            # If still no title but we have company, try to find the title nearby
            if not job["title"] and job["company"]:
                # Look for capitalized phrases before or after company name
                company_idx = paragraph.find(job["company"])
                if company_idx != -1:
                    # Check before company
                    before_text = paragraph[:company_idx].strip()
                    title_match = re.search(r'([A-Z][a-zA-Z\s]{3,30})$', before_text)
                    if title_match:
                        job["title"] = title_match.group(1).strip()
                    else:
                        # Check after company
                        after_text = paragraph[company_idx + len(job["company"]):].strip()
                        title_match = re.search(r'^([A-Z][a-zA-Z\s]{3,30})', after_text)
                        if title_match:
                            job["title"] = title_match.group(1).strip()
            
            # Clean up title if needed
            if job["title"]:
                # Remove any duration that might be in the title
                for pattern in date_patterns:
                    job["title"] = re.sub(pattern, "", job["title"]).strip()
                
                # Remove company name from title if it got included
                if job["company"]:
                    job["title"] = job["title"].replace(job["company"], "").strip()
                
                # Remove any trailing punctuation
                job["title"] = re.sub(r'[,;:]$', '', job["title"]).strip()
            
            # If we found meaningful information, add this job
            if job["duration"] or job["company"] or job["title"]:
                # Clean up the description
                # Remove bullet points and other formatting characters
                desc_lines = job["description"].split('\n')
                cleaned_lines = []
                
                for line in desc_lines:
                    # Remove bullet points and numbering
                    line = re.sub(r'^(?:•|\*|-|\d+\.|\–)\s+', '', line.strip())
                    if line:
                        cleaned_lines.append(line)
                
                job["description"] = '\n'.join(cleaned_lines)
                
                # Clean up incomplete sentences at beginning or end
                job["description"] = re.sub(r'^[,;:]\s+', '', job["description"])
                job["description"] = re.sub(r'[,;:]$', '.', job["description"])
                
                work_experiences.append(job)
    
    # If we couldn't extract structured work experience or if spaCy isn't available, 
    # try a regex-based approach
    if not work_experiences or not SPACY_AVAILABLE:
        # Find blocks of text that might be job entries
        blocks = []
        current_block = []
        
        # If we have an experience section, use that
        if experience_start_idx is not None and experience_end_idx is not None:
            lines_to_process = lines[experience_start_idx:experience_end_idx]
        else:
            # Otherwise try to find date ranges throughout the document
            lines_to_process = lines
        
        # Process lines to identify job blocks
        for line in lines_to_process:
            line = line.strip()
            if not line:
                if current_block:
                    blocks.append('\n'.join(current_block))
                    current_block = []
            else:
                # Check if this line might start a new job entry
                date_match = re.search(date_pattern, line, re.IGNORECASE)
                company_match = re.search(r'(?:[A-Z][a-z]*\s*(?:Inc|LLC|Ltd|Co|Corp|Corporation|Company|Group))|(?:[A-Z][A-Z\s]+)', line)
                
                if (date_match or company_match) and current_block:
                    # If this looks like a new job entry and we have content, save the current block
                    blocks.append('\n'.join(current_block))
                    current_block = [line]
                else:
                    current_block.append(line)
        
        # Add the last block if there is one
        if current_block:
            blocks.append('\n'.join(current_block))
        
        # Process each block to extract job information
        for block in blocks:
            # Initialize job entry
            job = {
                "title": "",
                "company": "",
                "duration": "",
                "location": "",
                "description": block
            }
            
            # Extract date range
            date_match = re.search(date_pattern, block, re.IGNORECASE)
            if date_match:
                job["duration"] = date_match.group(0).strip()
                # Remove date from description for cleaner text
                job["description"] = job["description"].replace(date_match.group(0), "").strip()
            
            # Extract company name
            company_pattern = r'(?:[A-Z][a-z]*\s*(?:Inc|LLC|Ltd|Co|Corp|Corporation|Company|Group))|(?:[A-Z][A-Z\s]+)'
            company_match = re.search(company_pattern, block)
            if company_match:
                job["company"] = company_match.group(0).strip()
            
            # Extract job title
            title_pattern = r'([A-Z][a-zA-Z\s]{3,30}(?:Engineer|Developer|Manager|Director|Analyst|Designer|Consultant|Specialist|Officer|Administrator|Coordinator|Lead|Head|Chief|Architect))'
            title_match = re.search(title_pattern, block)
            if title_match:
                job["title"] = title_match.group(0).strip()
                # Ensure title doesn't include company name
                if job["company"] and job["company"] in job["title"]:
                    job["title"] = job["title"].replace(job["company"], "").strip()
            
            # Extract location
            location_pattern = r'(?:in|at)\s+([A-Za-z\s,]+)(?:$|[,.])'
            location_match = re.search(location_pattern, block)
            if location_match:
                job["location"] = location_match.group(1).strip()
            
            # Clean up the description
            if job["duration"] or job["company"] or job["title"]:
                # Get the first paragraph or line as a summary if the description is long
                desc_lines = job["description"].split('\n')
                cleaned_lines = []
                
                for line in desc_lines:
                    # Remove bullet points
                    line = re.sub(r'^(?:•|\*|-|\d+\.|\–)\s+', '', line.strip())
                    if line:
                        cleaned_lines.append(line)
                
                job["description"] = '\n'.join(cleaned_lines)
                
                # Add to work experiences
                work_experiences.append(job)
    
    # Sort work experiences by date (if available)
    def extract_year(duration):
        # Try to extract the most recent year from the duration string
        year_match = re.search(r'(\d{4})', duration)
        if year_match:
            return int(year_match.group(1))
        if re.search(r'present|current|now|ongoing', duration, re.IGNORECASE):
            return 9999  # Current job should be first
        return 0  # Default if no year found
    
    work_experiences.sort(key=lambda job: extract_year(job.get("duration", "")), reverse=True)
    
    return work_experiences
def parse_resume(file_path):
    """Main function to parse resume and extract information"""
    logger.info(f"Parsing resume: {file_path}")
    
    try:
        # Extract text from the file
        text = extract_text(file_path)
        
        if not text or "Error extracting text" in text:
            return {"error": text}
        
        # Extract information
        email = extract_email(text)
        phone = extract_phone(text)
        skills = extract_skills(text)
        certifications = extract_certifications(text)
        work_experience = extract_work_experience(text)
        
        # Assemble the result
        result = {
            "file_name": Path(file_path).name,
            "text_length": len(text),
            "email": email,
            "phone": phone,
            "skills": skills,
            "certifications": certifications,
            "work_experience": work_experience,
            "text_sample": text[:300] + "..." if len(text) > 300 else text
        }
        logger.info(f"Successfully parsed resume: {Path(file_path).name}")
        return result
    except Exception as e:
        logger.error(f"Error parsing resume: {str(e)}")
        return {"error": f"Error parsing resume: {str(e)}"}

if __name__ == "__main__":
    # Test the parser with a sample file
    import sys
    if len(sys.argv) > 1:
        result = parse_resume(sys.argv[1])
        print(result)
    else:
        print("Please provide a file path as a command line argument.")